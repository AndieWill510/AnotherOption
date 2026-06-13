#!/usr/bin/env python3
"""Build the Tiny Door book manuscript as a styled DOCX.

This script assembles Options/001 through Options/101 in the order defined by
Options/INDEX.md, strips repository metadata/navigation footers, and produces a
6x9 book-style DOCX in dist/.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OPTIONS_DIR = ROOT / "Options"
INDEX_PATH = OPTIONS_DIR / "INDEX.md"
DIST_DIR = ROOT / "dist"
OUTPUT = DIST_DIR / "Tiny_Doors_and_Pillow_Forts_You_Have_Another_Option.docx"

BOOK_TITLE = "Tiny Doors and Pillow Forts"
BOOK_SUBTITLE = "You Have Another Option"
AUTHOR = "Kevin \"Andie\" Williams"

FONT_BODY = "Georgia"
FONT_SANS = "Aptos"

INDEX_RE = re.compile(r"^(\d{3})\. \[(.+?)\]\((.+?\.md)\)\s*$")
HTML_COMMENT_RE = re.compile(r"^<!--.*?-->\s*", re.S)
NAV_RE = re.compile(r"\n?---\n\n<!-- TINYDOOR_NAV_START -->.*?<!-- TINYDOOR_NAV_END -->\s*\Z", re.S)
MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def parse_index() -> list[tuple[str, str, str]]:
    entries: list[tuple[str, str, str]] = []
    for line in INDEX_PATH.read_text(encoding="utf-8").splitlines():
        match = INDEX_RE.match(line)
        if match:
            entries.append(match.groups())
    if not entries:
        raise RuntimeError("No entries found in Options/INDEX.md")
    return entries


def clean_option_markdown(raw: str) -> tuple[str, list[str]]:
    raw = NAV_RE.sub("", raw).rstrip()
    raw = HTML_COMMENT_RE.sub("", raw).lstrip()
    lines = raw.splitlines()

    title = "Untitled"
    body_start = 0
    for idx, line in enumerate(lines):
        if line.startswith("# "):
            title = line[2:].strip()
        if line.strip() == "---":
            body_start = idx + 1
            break

    body_lines = lines[body_start:]
    return title, body_lines


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.gutter = Inches(0.15)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.8)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)

    title_style = styles["Title"]
    title_style.font.name = FONT_BODY
    title_style.font.size = Pt(28)
    title_style.font.bold = True

    subtitle_style = styles.add_style("Book Subtitle", 1)
    subtitle_style.font.name = FONT_BODY
    subtitle_style.font.size = Pt(16)
    subtitle_style.font.italic = True
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(18)

    author_style = styles.add_style("Book Author", 1)
    author_style.font.name = FONT_BODY
    author_style.font.size = Pt(12)
    author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    chapter_style = styles["Heading 1"]
    chapter_style.font.name = FONT_BODY
    chapter_style.font.size = Pt(16)
    chapter_style.font.bold = True
    chapter_style.paragraph_format.space_before = Pt(0)
    chapter_style.paragraph_format.space_after = Pt(18)
    chapter_style.paragraph_format.keep_with_next = True

    subhead_style = styles["Heading 2"]
    subhead_style.font.name = FONT_BODY
    subhead_style.font.size = Pt(11)
    subhead_style.font.bold = True
    subhead_style.paragraph_format.space_before = Pt(10)
    subhead_style.paragraph_format.space_after = Pt(5)
    subhead_style.paragraph_format.keep_with_next = True

    quote_style = styles.add_style("Tiny Door Quote", 1)
    quote_style.font.name = FONT_BODY
    quote_style.font.size = Pt(10.8)
    quote_style.font.italic = True
    quote_style.paragraph_format.left_indent = Inches(0.28)
    quote_style.paragraph_format.right_indent = Inches(0.18)
    quote_style.paragraph_format.space_before = Pt(4)
    quote_style.paragraph_format.space_after = Pt(8)

    toc_style = styles.add_style("Contents Entry", 1)
    toc_style.font.name = FONT_BODY
    toc_style.font.size = Pt(9.3)
    toc_style.paragraph_format.space_after = Pt(2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_text_with_markdown(paragraph, text: str) -> None:
    # Remove markdown links but keep the visible label.
    text = MARKDOWN_LINK_RE.sub(r"\1", text)

    # Simple inline **bold** / *italic* parser. This deliberately avoids cleverness.
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_body_lines(doc: Document, lines: Iterable[str]) -> None:
    pending_blank = False
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            pending_blank = True
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:].strip(), style="Heading 2")
            pending_blank = False
            continue

        if stripped.startswith(">"):
            quote = stripped.lstrip("> ").strip()
            p = doc.add_paragraph(style="Tiny Door Quote")
            add_text_with_markdown(p, quote)
            pending_blank = False
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_text_with_markdown(p, stripped[2:].strip())
            pending_blank = False
            continue

        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if ordered:
            p = doc.add_paragraph(style="List Number")
            add_text_with_markdown(p, ordered.group(1).strip())
            pending_blank = False
            continue

        p = doc.add_paragraph()
        if pending_blank:
            p.paragraph_format.space_before = Pt(2)
        add_text_with_markdown(p, stripped)
        pending_blank = False


def add_title_page(doc: Document) -> None:
    for _ in range(7):
        doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(BOOK_TITLE)

    subtitle = doc.add_paragraph(style="Book Subtitle")
    subtitle.add_run(BOOK_SUBTITLE)

    byline = doc.add_paragraph(style="Book Author")
    byline.add_run(f"by {AUTHOR}")


def add_copyright_page(doc: Document) -> None:
    add_page_break(doc)
    p = doc.add_paragraph()
    p.add_run(f"{BOOK_TITLE}: {BOOK_SUBTITLE}").bold = True
    doc.add_paragraph(f"Copyright 2026 {AUTHOR}")
    doc.add_paragraph("Licensed under the Apache License, Version 2.0.")
    doc.add_paragraph("This manuscript discusses trauma, suicide prevention, burnout, dysregulation, sensory overload, abuse, and institutional harm. It is not medical, legal, or crisis advice.")
    doc.add_paragraph("A tiny door is not the whole answer. It is one safe next passage.")


def add_contents(doc: Document, entries: list[tuple[str, str, str]]) -> None:
    add_page_break(doc)
    doc.add_paragraph("Contents", style="Heading 1")
    for number, title, _filename in entries:
        p = doc.add_paragraph(style="Contents Entry")
        p.add_run(f"{number} — {title}")


def build_book() -> None:
    entries = parse_index()
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_copyright_page(doc)
    add_contents(doc, entries)

    for number, index_title, filename in entries:
        add_page_break(doc)
        raw = (OPTIONS_DIR / filename).read_text(encoding="utf-8")
        title, body_lines = clean_option_markdown(raw)
        # Prefer the authored H1, but keep index title as fallback.
        title = title or f"{number} — {index_title}"
        doc.add_paragraph(title, style="Heading 1")
        add_body_lines(doc, body_lines)

    DIST_DIR.mkdir(exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    build_book()
